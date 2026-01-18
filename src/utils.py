import zipfile
import os
import tempfile
import fitz  # PyMuPDF
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH



def extract_paper_sections(text: str) -> dict:
    """
    Deterministically split a scientific paper into sections.
    Returns dict with explicit None for missing sections.
    """

    title_snippet = text[:500].strip() if text else None

    # Normalize text for heading detection
    normalized = re.sub(r'\s+', ' ', text)
    lowered = normalized.lower()

    CANONICAL_SECTIONS = {
        "abstract": ["abstract"],
        "introduction": ["introduction", "background"],
        "methods": ["methods", "materials and methods", "methodology"],
        "results": ["results", "findings"],
        "discussion": ["discussion"],
        "conclusion": ["conclusion", "conclusions"],
        "funding": ["funding", "funding statement", "sources of funding"],
        "conflicts_of_interest": [
            "conflicts of interest",
            "conflict of interest",
            "competing interests",
            "disclosure"
        ]
    }
    
    # Find all candidate headings with positions
    heading_positions = []

    for section, variants in CANONICAL_SECTIONS.items():
        for v in variants:
            pattern = rf'\b{re.escape(v)}\b'
            for m in re.finditer(pattern, lowered):
                heading_positions.append((m.start(), section))

    # Sort by position in document
    heading_positions.sort(key=lambda x: x[0])

    # Initialize output with None
    sections = {k: None for k in CANONICAL_SECTIONS.keys()}
    sections["title"] = title_snippet

    # Edge case: no headings found
    if not heading_positions:
        return sections

    # Slice text between headings
    for i, (start_idx, section_name) in enumerate(heading_positions):
        end_idx = heading_positions[i + 1][0] if i + 1 < len(heading_positions) else len(normalized)
        section_text = normalized[start_idx:end_idx].strip()

        # Only keep the *first* occurrence of each section
        if sections[section_name] is None:
            sections[section_name] = section_text

    return sections


def clean_text_for_llm(raw_text: str) -> str:
    """Normalize text for LLM JSON generation."""
    # Replace common PDF Unicode issues
    replacements = {
        '“': '"', '”': '"', '‘': "'", '’': "'",
        '—': '-', '–': '-', ' ': ' ', '\u00a0': ' ',  # Non-breaking space
        '\t': ' ', '\n': ' ', '\r': ' '
    }
    for bad, good in replacements.items():
        raw_text = raw_text.replace(bad, good)
    
    # Remove all control characters except basic whitespace
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', ' ', raw_text)
    return ' '.join(cleaned.split())  # Normalize whitespace

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extracts text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return clean_text_for_llm(text)
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def process_zip_file(uploaded_file) -> list:
    """
    Unzips file and returns a list of dictionaries.
    Returns: [{'filename': str, 'content': str}, ...]
    """
    papers_data = []
    
    with tempfile.TemporaryDirectory() as temp_dir:
        zip_path = os.path.join(temp_dir, "upload.zip")
        extract_dir = os.path.join(temp_dir, "extracted")
        os.makedirs(extract_dir, exist_ok=True)

        with open(zip_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
            
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    # Filter for PDFs and ignore hidden files
                    if file.lower().endswith(".pdf") and not file.startswith("._"):
                        full_path = os.path.join(root, file)
                        text_content = extract_text_from_pdf(full_path)
                    if not file.lower().endswith(".pdf"):
                        continue
                        
                    sections = extract_paper_sections(text_content)

                    papers_data.append({
                        "filename": file,
                        "sections": sections
                    })
    return papers_data


def generate_docx_report(results: list) -> BytesIO:
    """
    Generates a Word document from the list of analysis results.
    Returns a BytesIO object containing the .docx file.
    """
    doc = Document()
    
    # --- Title Page ---
    title = doc.add_heading('Nutrition Science Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Total Papers Analyzed: {len(results)}")
    doc.add_page_break()

    # --- Loop through each paper ---
    for i, res in enumerate(results):
        # SAFER title logic - check title first, fallback to index/fallback
        if hasattr(res, 'title') and res.title and res.title.strip():
            paper_title = res.title[:100]  # Truncate long titles
        else:
            paper_title = f"Paper {i+1} ({get_filename_fallback(res)})"
            
        header = doc.add_heading(paper_title, level=1)
        
        # Sub-header details using exact schema fields
        p = doc.add_paragraph()
        p.add_run(f"Type: {getattr(res, 'paper_type', 'N/A')} | ").bold = True
        p.add_run(f"Evidence Level: {getattr(res, 'evidence_level', 'N/A')} | ").bold = True
        
        # Safe trust score access
        trust_score = getattr(res, 'trust_score', 0)
        run = p.add_run(f"Trust Score: {trust_score}/10")
        run.bold = True
        if isinstance(trust_score, (int, float)) and trust_score < 5:
            run.font.color.rgb = RGBColor(255, 0, 0)
        
        # 1. Conflicts of Interest
        doc.add_heading('💰 Conflicts of Interest & Funding', level=2)
        has_coi = getattr(res, 'has_conflict_of_interest', None)

        if has_coi:
            p = doc.add_paragraph()
            warning = p.add_run("⚠️ WARNING: CONFLICT DETECTED")
            warning.font.color.rgb = RGBColor(255, 0, 0)
            warning.bold = True
            doc.add_paragraph(f"Source: {getattr(res, 'funding_source', 'N/A')}")
            doc.add_paragraph(f"Notes: {getattr(res, 'coi_notes', 'N/A')}")
        elif has_coi is False:
            doc.add_paragraph("No conflicts of interest declared by the authors.")
        else:
            doc.add_paragraph("Conflict of interest statement NOT REPORTED.")

        # 2. Methodology
        doc.add_heading('🔬 Methodology Audit', level=2)
        doc.add_paragraph(f"Control Group: {getattr(res, 'control_group_quality', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Intervention: {getattr(res, 'intervention_details', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Confounders: {getattr(res, 'confounding_factors', 'N/A')}", style='List Bullet')

        # 3. Statistics
        doc.add_heading('📈 Statistical Analysis', level=2)
        doc.add_paragraph(f"Primary Outcome: {getattr(res, 'primary_outcome', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Risk Reported: {getattr(res, 'risk_type_reported', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Endpoints: {getattr(res, 'endpoints', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Significance: {getattr(res, 'statistical_significance', 'N/A')}", style='List Bullet')

        # 4. Conclusions
        doc.add_heading('🏁 Conclusions & Verdict', level=2)
        doc.add_paragraph(f"Authors' Conclusion: {getattr(res, 'conclusion_summary', 'N/A')}")
        doc.add_paragraph(getattr(res, 'final_verdict', 'N/A'), style='List Bullet')
        
        # Separator between papers
        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def get_filename_fallback(res) -> str:
    """Safe filename fallback - your app.py likely passes this."""
    # Common patterns where filename might be stored
    try:
        # Option 1: Check if results have filename from process_zip_file
        if hasattr(res, 'filename') and res.filename:
            return res.filename.replace('.pdf', '').title()
        # Option 2: Use generic fallback
        return "Untitled Analysis"
    except:
        return "Untitled Analysis"
